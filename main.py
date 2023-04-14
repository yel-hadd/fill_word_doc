from os import system
import pandas as pd
import docx


def generate_docs(excel_path, word_path):
    system("rm -rf output/*")
    df = pd.read_excel(excel_path)

    columns = list(df.columns)

    for index, row in df.iterrows():
        doc = docx.Document(word_path)
        for column in columns:
            pattern = f"<<{column}>>"
            try:
                outpath = f"./output/{row[columns[0]]}_{row[columns[1]].replace(' ', '_').lower()}.docx"
            except:
                outpath = f"{index}.docx"
            for paragraph in doc.paragraphs:
                if pattern in paragraph.text:
                    paragraph.text = paragraph.text.replace(pattern, str(row[column]))
            for table in doc.tables:
                for r in table.rows:
                    for cell in r.cells:
                        if pattern in cell.text:
                            cell.text = cell.text.replace(pattern, str(row[column]))
        doc.save(outpath)
        print(outpath + " has been written")


generate_docs("./excel/path/here", "./word/path/here")
